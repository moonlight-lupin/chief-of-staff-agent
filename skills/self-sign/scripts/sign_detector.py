#!/usr/bin/env python3
"""Detect signature and date locations in PDF or DOCX documents.

The detector is intentionally read-only: it finds candidate locations and returns
structured metadata so the Self-Sign skill can ask the user which blocks to sign.

Supported formats:
  - PDF via PyMuPDF (imported lazily, with a graceful error if unavailable)
  - DOCX via python-docx

Example:
    python sign_detector.py agreement.pdf --format json --company "Acme Pte Ltd" \
        --alias "Service Provider" --alias Consultant
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable, Iterator, Optional


DEFAULT_SIGNATURE_PATTERNS = [
    r"Signature\s*:",
    r"Signed\s+by\s*:",
    r"Sign\s+here\s*:",
    r"Authori[sz]ed\s+Signator(?:y|ies)\s*:",
    r"For\s+and\s+on\s+behalf\s+of",
    r"Executed\s+by",
    r"Name\s*:\s*_{5,}",
    r"_{10,}",
]

DEFAULT_DATE_PATTERNS = [
    r"Date\s*:",
    r"Date\s+Signed\s*:",
    r"Dated\s+this",
    r"_{10,}\s*(?:Date|Dated)",
]

DEFAULT_PARTY_PATTERNS = [
    r"For\s+and\s+on\s+behalf\s+of(?:\s+the)?\s+[^\n:]{0,120}",
    r"Signed\s+by(?:\s+the)?\s+[^\n:]{0,120}",
    r"Party\s+[AB]\b",
    r"First\s+Party\b",
    r"Second\s+Party\b",
    r"The\s+Client\b",
    r"The\s+Service\s+Provider\b",
    r"The\s+Consultant\b",
    r"Consultant\b",
    r"Contractor\b",
    r"Service\s+Provider\b",
    r"Client\b",
    r"Lessor\b",
    r"Lessee\b",
    r"Buyer\b",
    r"Seller\b",
    r"Employer\b",
    r"Employee\b",
]

PROFESSIONAL_SERVICE_ALIASES = ["service provider", "consultant", "contractor", "the company"]


@dataclass
class SignatureLocation:
    """A detected signature or date location.

    Attributes:
        page: 1-based PDF page number, or None for DOCX.
        paragraph: DOCX paragraph index, or None for PDF.
        coordinates: PDF bounding box [x0, y0, x1, y1], or None for DOCX.
        matched_text: Text that triggered detection.
        party_context: Nearby party label/context, if found.
        confidence: Float in [0, 1] estimating whether this is a true target.
        location_type: "signature" or "date".
        matched_party: "self", "other", or "unknown" based on company/alias matching.
        source: "pdf" or "docx".
    """

    page: Optional[int]
    paragraph: Optional[int]
    coordinates: Optional[list[float]]
    matched_text: str
    party_context: Optional[str]
    confidence: float
    location_type: str = "signature"
    matched_party: str = "unknown"
    source: str = "unknown"


@dataclass
class TextLine:
    text: str
    bbox: tuple[float, float, float, float]
    page: int
    order: int


def _compile(patterns: Iterable[str]) -> list[re.Pattern[str]]:
    return [re.compile(p, re.IGNORECASE) for p in patterns]


def _normalize(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _shorten(text: str, limit: int = 220) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _pattern_type(text: str, sig_patterns: list[re.Pattern[str]], date_patterns: list[re.Pattern[str]]) -> str:
    if any(p.search(text) for p in date_patterns) and not any(p.search(text) for p in sig_patterns):
        return "date"
    if re.search(r"\bdate\b", text, re.IGNORECASE) and not re.search(
        r"signature|signed|signatory|behalf", text, re.IGNORECASE
    ):
        return "date"
    return "signature"


def _has_location_pattern(
    text: str,
    sig_patterns: list[re.Pattern[str]],
    date_patterns: list[re.Pattern[str]],
) -> bool:
    return any(p.search(text) for p in sig_patterns) or any(p.search(text) for p in date_patterns)


def _find_party_context_in_text(text: str, party_patterns: list[re.Pattern[str]]) -> Optional[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    candidates: list[str] = []
    for line in lines:
        for pattern in party_patterns:
            match = pattern.search(line)
            if match:
                candidates.append(_shorten(line))
                break
    if candidates:
        return candidates[-1]
    return _shorten(text) if text.strip() else None


def _classify_party(party_context: Optional[str], company: Optional[str], aliases: list[str]) -> str:
    if not party_context:
        return "unknown"
    context_norm = _normalize(party_context)
    positives = [_normalize(a) for a in aliases if a]
    if company:
        positives.append(_normalize(company))
    for needle in positives:
        if needle and needle in context_norm:
            return "self"
    negative_terms = [
        "the client",
        "client",
        "buyer",
        "lessee",
        "employee",
        "customer",
    ]
    if any(term in context_norm for term in negative_terms) and not any(
        alias in context_norm for alias in PROFESSIONAL_SERVICE_ALIASES
    ):
        return "other"
    return "unknown"


def _confidence(
    matched_text: str,
    party_context: Optional[str],
    matched_party: str,
    location_type: str,
) -> float:
    score = 0.45
    if re.search(r"_{10,}", matched_text):
        score += 0.18
    if re.search(r"signature|signed|signatory|behalf|executed", matched_text, re.IGNORECASE):
        score += 0.18
    if location_type == "date" and re.search(r"date|dated", matched_text, re.IGNORECASE):
        score += 0.18
    if party_context:
        score += 0.12
    if matched_party == "self":
        score += 0.07
    elif matched_party == "other":
        score -= 0.05
    return round(max(0.05, min(score, 0.99)), 2)


def _dedupe(locations: list[SignatureLocation]) -> list[SignatureLocation]:
    seen: set[tuple[Any, ...]] = set()
    result: list[SignatureLocation] = []
    for loc in sorted(
        locations,
        key=lambda l: (
            l.page or 0,
            l.paragraph if l.paragraph is not None else -1,
            l.coordinates[1] if l.coordinates else -1,
            l.coordinates[0] if l.coordinates else -1,
            l.location_type,
        ),
    ):
        coord_key = tuple(round(x, 1) for x in loc.coordinates) if loc.coordinates else None
        key = (loc.page, loc.paragraph, coord_key, _normalize(loc.matched_text)[:80], loc.location_type)
        if key in seen:
            continue
        seen.add(key)
        result.append(loc)
    return result


def _extract_pdf_lines(pdf_path: Path) -> list[TextLine]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "PyMuPDF is required for PDF signature detection. Install it with `python3 -m pip install pymupdf`."
        ) from exc

    lines: list[TextLine] = []
    try:
        doc = fitz.open(str(pdf_path))
    except Exception as exc:  # pragma: no cover - depends on malformed files
        raise RuntimeError(f"Could not open PDF {pdf_path}: {exc}") from exc

    try:
        order = 0
        for page_index, page in enumerate(doc, start=1):
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text = "".join(span.get("text", "") for span in spans).strip()
                    if not text:
                        continue
                    bbox = tuple(float(v) for v in line.get("bbox", block.get("bbox", (0, 0, 0, 0))))
                    lines.append(TextLine(text=text, bbox=bbox, page=page_index, order=order))
                    order += 1
    finally:
        doc.close()
    return lines


def _scan_upward_pdf(
    lines: list[TextLine],
    hit: TextLine,
    party_patterns: list[re.Pattern[str]],
    max_points: float = 260.0,
    max_lines: int = 12,
) -> Optional[str]:
    candidates: list[tuple[float, TextLine]] = []
    x0, y0, x1, _ = hit.bbox
    for line in lines:
        if line.page != hit.page or line.order >= hit.order:
            continue
        lx0, ly0, lx1, ly1 = line.bbox
        vertical_gap = y0 - ly1
        if vertical_gap < -4 or vertical_gap > max_points:
            continue
        horizontal_overlap = max(0.0, min(x1, lx1) - max(x0, lx0))
        same_column = horizontal_overlap > 0 or abs(lx0 - x0) < 140
        if same_column:
            candidates.append((vertical_gap, line))
    candidates = sorted(candidates, key=lambda item: item[0])[:max_lines]
    for _, line in candidates:
        if _find_party_context_in_text(line.text, party_patterns):
            return _shorten(line.text)
    if candidates:
        context = " | ".join(line.text for _, line in reversed(candidates[-3:]))
        return _find_party_context_in_text(context, party_patterns)
    return None


def detect_pdf(
    path: Path,
    company: Optional[str] = None,
    aliases: Optional[list[str]] = None,
    signature_patterns: Optional[list[str]] = None,
    date_patterns: Optional[list[str]] = None,
    party_patterns: Optional[list[str]] = None,
) -> list[SignatureLocation]:
    """Scan a PDF for signature/date locations and return structured candidates."""
    aliases = aliases or []
    sig_re = _compile(signature_patterns or DEFAULT_SIGNATURE_PATTERNS)
    date_re = _compile(date_patterns or DEFAULT_DATE_PATTERNS)
    party_re = _compile(party_patterns or DEFAULT_PARTY_PATTERNS)
    lines = _extract_pdf_lines(path)
    locations: list[SignatureLocation] = []

    for line in lines:
        if not _has_location_pattern(line.text, sig_re, date_re):
            continue
        location_type = _pattern_type(line.text, sig_re, date_re)
        party_context = _scan_upward_pdf(lines, line, party_re)
        if not party_context:
            party_context = _find_party_context_in_text(line.text, party_re)
        matched_party = _classify_party(party_context, company, aliases)
        locations.append(
            SignatureLocation(
                page=line.page,
                paragraph=None,
                coordinates=[round(v, 2) for v in line.bbox],
                matched_text=_shorten(line.text),
                party_context=party_context,
                confidence=_confidence(line.text, party_context, matched_party, location_type),
                location_type=location_type,
                matched_party=matched_party,
                source="pdf",
            )
        )
    return _dedupe(locations)


def _iter_docx_paragraph_text(document: Any) -> Iterator[tuple[int, str]]:
    index = 0
    for paragraph in document.paragraphs:
        yield index, paragraph.text
        index += 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield index, paragraph.text
                    index += 1


def detect_docx(
    path: Path,
    company: Optional[str] = None,
    aliases: Optional[list[str]] = None,
    signature_patterns: Optional[list[str]] = None,
    date_patterns: Optional[list[str]] = None,
    party_patterns: Optional[list[str]] = None,
) -> list[SignatureLocation]:
    """Scan a DOCX for signature/date locations and return structured candidates."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX signature detection. Install it with `python3 -m pip install python-docx`."
        ) from exc

    aliases = aliases or []
    sig_re = _compile(signature_patterns or DEFAULT_SIGNATURE_PATTERNS)
    date_re = _compile(date_patterns or DEFAULT_DATE_PATTERNS)
    party_re = _compile(party_patterns or DEFAULT_PARTY_PATTERNS)

    try:
        document = Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"Could not open DOCX {path}: {exc}") from exc

    paragraphs = [(idx, text) for idx, text in _iter_docx_paragraph_text(document)]
    locations: list[SignatureLocation] = []
    for pos, (paragraph_idx, text) in enumerate(paragraphs):
        if not text or not _has_location_pattern(text, sig_re, date_re):
            continue
        location_type = _pattern_type(text, sig_re, date_re)
        surrounding = "\n".join(t for _, t in paragraphs[max(0, pos - 5) : pos + 1] if t.strip())
        party_context = _find_party_context_in_text(surrounding, party_re)
        matched_party = _classify_party(party_context, company, aliases)
        locations.append(
            SignatureLocation(
                page=None,
                paragraph=paragraph_idx,
                coordinates=None,
                matched_text=_shorten(text),
                party_context=party_context,
                confidence=_confidence(text, party_context, matched_party, location_type),
                location_type=location_type,
                matched_party=matched_party,
                source="docx",
            )
        )
    return _dedupe(locations)


def detect_signatures(
    path: str | Path,
    doc_type: str = "auto",
    company: Optional[str] = None,
    aliases: Optional[list[str]] = None,
) -> list[SignatureLocation]:
    """Detect signature locations in a PDF or DOCX file."""
    document_path = Path(path).expanduser().resolve()
    if not document_path.exists():
        raise FileNotFoundError(f"Document not found: {document_path}")
    suffix = document_path.suffix.lower()
    if doc_type == "auto":
        if suffix == ".pdf":
            doc_type = "pdf"
        elif suffix == ".docx":
            doc_type = "docx"
        else:
            raise ValueError("Could not infer document type. Use --doc-type pdf or --doc-type docx.")
    if doc_type == "pdf":
        return detect_pdf(document_path, company=company, aliases=aliases)
    if doc_type == "docx":
        return detect_docx(document_path, company=company, aliases=aliases)
    raise ValueError(f"Unsupported document type: {doc_type}")


def _load_aliases_from_file(path: Optional[str]) -> list[str]:
    if not path:
        return []
    config_path = Path(path).expanduser()
    try:
        import yaml  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyYAML is required to read --config files.") from exc
    data = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}
    self_sign = data.get("self_sign", {}) if isinstance(data, dict) else {}
    aliases = self_sign.get("party_aliases", []) if isinstance(self_sign, dict) else []
    if not isinstance(aliases, list):
        return []
    return [str(a) for a in aliases]


def _company_from_file(path: Optional[str]) -> Optional[str]:
    if not path:
        return None
    config_path = Path(path).expanduser()
    try:
        import yaml  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyYAML is required to read --config files.") from exc
    data = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}
    company = data.get("company", {}) if isinstance(data, dict) else {}
    name = company.get("name") if isinstance(company, dict) else None
    return str(name) if name else None


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Detect signature/date locations in PDF or DOCX documents.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("document", help="Path to PDF or DOCX document")
    parser.add_argument("--doc-type", choices=["auto", "pdf", "docx"], default="auto", help="Document type")
    parser.add_argument("--format", choices=["json", "text"], default="json", help="Output format")
    parser.add_argument("--company", help="Company name used to classify self-owned signature blocks")
    parser.add_argument("--alias", action="append", default=[], help="Party alias that should match the user's company; repeatable")
    parser.add_argument("--config", help="Optional company.yaml to load company.name and self_sign.party_aliases")
    parser.add_argument("--include-professional-defaults", action="store_true", help="Add Service Provider/Consultant/Contractor aliases")
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        company = args.company or _company_from_file(args.config)
        aliases = list(args.alias) + _load_aliases_from_file(args.config)
        if args.include_professional_defaults:
            aliases.extend(PROFESSIONAL_SERVICE_ALIASES)
        locations = detect_signatures(args.document, doc_type=args.doc_type, company=company, aliases=aliases)
    except (FileNotFoundError, ValueError, RuntimeError) as exc:
        print(f"sign_detector error: {exc}", file=sys.stderr)
        return 2

    payload = [asdict(loc) for loc in locations]
    if args.format == "json":
        print(json.dumps(payload, indent=2, ensure_ascii=False))
    else:
        if not payload:
            print("No signature locations found.")
        for i, loc in enumerate(locations, start=1):
            where = f"page {loc.page}" if loc.page is not None else f"paragraph {loc.paragraph}"
            print(f"{i}. {where} [{loc.location_type}] confidence={loc.confidence} party={loc.matched_party}")
            print(f"   matched: {loc.matched_text}")
            if loc.party_context:
                print(f"   context: {loc.party_context}")
            if loc.coordinates:
                print(f"   bbox: {loc.coordinates}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
