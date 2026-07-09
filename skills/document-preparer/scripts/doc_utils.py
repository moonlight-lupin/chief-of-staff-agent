#!/usr/bin/env python3
"""DOCX template utilities for the Chief of Staff Document Preparer skill.

Functions are importable and the module also provides a small CLI:

    python doc_utils.py fill --template X.docx --output Y.docx --tokens '{"client_name":"Acme"}'
    python doc_utils.py extract --template X.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable, Optional

TOKEN_RE = re.compile(r"{{\s*([A-Za-z_][A-Za-z0-9_.-]*)\s*}}")


class DocUtilsError(RuntimeError):
    """Raised for user-facing document utility errors."""


def _require_docx() -> Any:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency presence varies
        raise DocUtilsError("python-docx is required. Install it with `python3 -m pip install python-docx`.") from exc
    return Document


def _load_document(path: str | Path) -> Any:
    document_path = Path(path).expanduser()
    if document_path.suffix.lower() != ".docx":
        raise DocUtilsError(f"Only .docx files are supported in Phase 1: {document_path}")
    if not document_path.exists():
        raise DocUtilsError(f"Document not found: {document_path}")
    Document = _require_docx()
    try:
        return Document(str(document_path))
    except Exception as exc:
        raise DocUtilsError(f"Could not open DOCX {document_path}: {exc}") from exc


def _iter_paragraphs(document: Any) -> Iterable[Any]:
    for paragraph in document.paragraphs:
        yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    """Replace paragraph text while preserving the first run's style where possible."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> int:
    original = paragraph.text
    if not original:
        return 0
    updated = original
    replacement_count = 0
    for needle, value in replacements.items():
        count = updated.count(needle)
        if count:
            updated = updated.replace(needle, value)
            replacement_count += count
    if updated != original:
        _set_paragraph_text(paragraph, updated)
    return replacement_count


def _token_replacements(tokens_dict: dict[str, Any]) -> dict[str, str]:
    replacements: dict[str, str] = {}
    for key, value in tokens_dict.items():
        token_name = str(key).strip()
        if token_name.startswith("{{") and token_name.endswith("}}"):
            needle = token_name
        else:
            needle = "{{" + token_name + "}}"
        replacements[needle] = "" if value is None else str(value)
        # Also support whitespace variants like {{ client_name }}.
        bare = token_name.strip("{} ")
        replacements[f"{{{{ {bare} }}}}"] = "" if value is None else str(value)
    return replacements


def fill_template(template_path: str | Path, tokens_dict: dict[str, Any], output_path: str | Path) -> dict[str, Any]:
    """Replace {{tokens}} in a DOCX template and save to output_path.

    Returns metadata with replacement count and any remaining tokens.
    """
    document = _load_document(template_path)
    replacements = _token_replacements(tokens_dict)
    total_replacements = 0
    for paragraph in _iter_paragraphs(document):
        total_replacements += _replace_in_paragraph(paragraph, replacements)
    destination = Path(output_path).expanduser()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    remaining = extract_tokens(destination)
    return {
        "output": str(destination),
        "replacements": total_replacements,
        "remaining_tokens": sorted(remaining),
    }


def extract_tokens(template_path: str | Path) -> set[str]:
    """Return all {{token}} placeholder names found in a DOCX file."""
    document = _load_document(template_path)
    tokens: set[str] = set()
    for paragraph in _iter_paragraphs(document):
        for match in TOKEN_RE.finditer(paragraph.text or ""):
            tokens.add(match.group(1))
    return tokens


def _normalize_mapping(token_mappings: dict[str, Any]) -> dict[str, str]:
    """Normalize mappings to literal_text -> {{token}} replacements.

    Accepts either {"literal text": "token_name"} or {"token_name": "literal text"}
    when the key already looks like a token placeholder.
    """
    replacements: dict[str, str] = {}
    for key, value in token_mappings.items():
        k = str(key)
        v = str(value)
        if k.startswith("{{") and k.endswith("}}"):
            replacements[v] = k
        elif v.startswith("{{") and v.endswith("}}"):
            replacements[k] = v
        elif re.fullmatch(r"[A-Za-z_][A-Za-z0-9_.-]*", v):
            replacements[k] = "{{" + v + "}}"
        elif re.fullmatch(r"[A-Za-z_][A-Za-z0-9_.-]*", k):
            replacements[v] = "{{" + k + "}}"
        else:
            raise DocUtilsError(
                f"Cannot infer token mapping direction for {key!r}: {value!r}. Use literal->token_name."
            )
    return replacements


def create_template_from_doc(
    doc_path: str | Path,
    token_mappings: dict[str, Any],
    output_path: str | Path,
) -> dict[str, Any]:
    """Create a reusable DOCX template by replacing literal text with {{tokens}}."""
    document = _load_document(doc_path)
    replacements = _normalize_mapping(token_mappings)
    total_replacements = 0
    missing_literals: list[str] = []
    all_text = "\n".join(paragraph.text for paragraph in _iter_paragraphs(document))
    for literal in replacements:
        if literal not in all_text:
            missing_literals.append(literal)
    for paragraph in _iter_paragraphs(document):
        total_replacements += _replace_in_paragraph(paragraph, replacements)
    destination = Path(output_path).expanduser()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return {
        "output": str(destination),
        "replacements": total_replacements,
        "tokens": sorted(extract_tokens(destination)),
        "missing_literals": missing_literals,
    }


def _load_yaml(path: Path) -> dict[str, Any]:
    try:
        import yaml  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise DocUtilsError("PyYAML is required for template registry operations.") from exc
    if not path.exists():
        return {"templates": []}
    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    if data is None:
        return {"templates": []}
    if not isinstance(data, dict):
        raise DocUtilsError(f"Template index must be a YAML mapping: {path}")
    data.setdefault("templates", [])
    if not isinstance(data["templates"], list):
        raise DocUtilsError("Template index key 'templates' must be a list.")
    return data


def _dump_yaml(path: Path, data: dict[str, Any]) -> None:
    try:
        import yaml  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise DocUtilsError("PyYAML is required for template registry operations.") from exc
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.safe_dump(data, sort_keys=False, allow_unicode=True), encoding="utf-8")


def register_template(
    name: str,
    file: str,
    tokens: list[str],
    category: str,
    index_path: str | Path,
    description: Optional[str] = None,
) -> dict[str, Any]:
    """Add or update a template entry in the YAML registry."""
    path = Path(index_path).expanduser()
    data = _load_yaml(path)
    templates = data["templates"]
    entry = {
        "name": name,
        "file": file,
        "tokens": sorted({str(t).strip("{} ") for t in tokens if str(t).strip("{} ")}),
        "category": category,
        "description": description or "",
        "last_used": None,
    }
    replaced = False
    for i, existing in enumerate(templates):
        if isinstance(existing, dict) and existing.get("name") == name:
            # Preserve last_used if present.
            entry["last_used"] = existing.get("last_used")
            templates[i] = entry
            replaced = True
            break
    if not replaced:
        templates.append(entry)
    templates.sort(key=lambda item: str(item.get("name", "")).lower() if isinstance(item, dict) else "")
    _dump_yaml(path, data)
    return {"index": str(path), "name": name, "updated": replaced, "entry": entry}


def _parse_json_mapping(raw: str, label: str) -> dict[str, Any]:
    try:
        value = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise DocUtilsError(f"Invalid JSON for {label}: {exc}") from exc
    if not isinstance(value, dict):
        raise DocUtilsError(f"{label} must be a JSON object.")
    return value


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="DOCX template utilities for Chief of Staff Document Preparer.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    fill = sub.add_parser("fill", help="Fill a DOCX template with token values")
    fill.add_argument("--template", required=True, help="Path to .docx template")
    fill.add_argument("--output", required=True, help="Output .docx path")
    fill.add_argument("--tokens", required=True, help="JSON object mapping token names to values")
    fill.add_argument("--format", choices=["json", "text"], default="json")

    extract = sub.add_parser("extract", help="List {{tokens}} in a DOCX template")
    extract.add_argument("--template", required=True, help="Path to .docx template")
    extract.add_argument("--format", choices=["json", "text"], default="text")

    templ = sub.add_parser("template", help="Create a template by replacing literal text with tokens")
    templ.add_argument("--doc", required=True, help="Source .docx document")
    templ.add_argument("--output", required=True, help="Output .docx template")
    templ.add_argument("--mappings", required=True, help="JSON object mapping literal text to token names")
    templ.add_argument("--format", choices=["json", "text"], default="json")

    reg = sub.add_parser("register", help="Add/update a template registry entry")
    reg.add_argument("--name", required=True)
    reg.add_argument("--file", required=True)
    reg.add_argument("--tokens", nargs="+", required=True)
    reg.add_argument("--category", required=True)
    reg.add_argument("--index", default="/root/.hermes/plugins/chief-of-staff/shared/templates/index.yaml")
    reg.add_argument("--description", default="")
    reg.add_argument("--format", choices=["json", "text"], default="json")
    return parser


def _print_payload(payload: Any, output_format: str) -> None:
    if output_format == "json":
        print(json.dumps(payload, indent=2, ensure_ascii=False))
    else:
        if isinstance(payload, (set, list, tuple)):
            for item in sorted(payload):
                print(item)
        elif isinstance(payload, dict):
            for key, value in payload.items():
                print(f"{key}: {value}")
        else:
            print(payload)


def main(argv: Optional[list[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        if args.command == "fill":
            payload = fill_template(args.template, _parse_json_mapping(args.tokens, "--tokens"), args.output)
            _print_payload(payload, args.format)
        elif args.command == "extract":
            payload = sorted(extract_tokens(args.template))
            _print_payload(payload, args.format)
        elif args.command == "template":
            payload = create_template_from_doc(args.doc, _parse_json_mapping(args.mappings, "--mappings"), args.output)
            _print_payload(payload, args.format)
        elif args.command == "register":
            payload = register_template(args.name, args.file, args.tokens, args.category, args.index, args.description)
            _print_payload(payload, args.format)
        else:  # pragma: no cover - argparse prevents this
            parser.error("Unknown command")
    except DocUtilsError as exc:
        print(f"doc_utils error: {exc}", file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
