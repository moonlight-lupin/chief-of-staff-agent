#!/usr/bin/env python3
"""Tests for doc_utils.py — template token extraction and filling."""

import sys
import json
import tempfile
from pathlib import Path

import pytest

DOC_SCRIPTS = Path(__file__).resolve().parents[1] / "skills" / "document-preparer" / "scripts"
if str(DOC_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(DOC_SCRIPTS))

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestExtractTokens:
    def test_extract_single_token(self, tmp_path):
        from doc_utils import extract_tokens
        docx_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("Dear {{client_name}},")
        doc.save(str(docx_path))
        tokens = extract_tokens(docx_path)
        assert "client_name" in tokens

    def test_extract_multiple_tokens(self, tmp_path):
        from doc_utils import extract_tokens
        docx_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("{{client_name}}")
        doc.add_paragraph("Date: {{date}}")
        doc.add_paragraph("Amount: {{amount}} {{currency}}")
        doc.save(str(docx_path))
        tokens = extract_tokens(docx_path)
        assert "client_name" in tokens
        assert "date" in tokens
        assert "amount" in tokens
        assert "currency" in tokens

    def test_extract_no_tokens(self, tmp_path):
        from doc_utils import extract_tokens
        docx_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("This is a plain document with no tokens.")
        doc.save(str(docx_path))
        tokens = extract_tokens(docx_path)
        assert len(tokens) == 0

    def test_extract_returns_set(self, tmp_path):
        from doc_utils import extract_tokens
        docx_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("{{a}} {{b}} {{a}}")
        doc.save(str(docx_path))
        tokens = extract_tokens(docx_path)
        assert isinstance(tokens, set)
        # Duplicates should be deduplicated
        assert len(tokens) == 2


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestFillTemplate:
    def test_fill_replaces_tokens(self, tmp_path):
        from doc_utils import fill_template, extract_tokens
        tmpl_path = tmp_path / "tmpl.docx"
        out_path = tmp_path / "out.docx"
        doc = Document()
        doc.add_paragraph("Dear {{client_name}},")
        doc.add_paragraph("Amount: {{amount}}")
        doc.save(str(tmpl_path))

        result = fill_template(tmpl_path, {
            "client_name": "Acme Corp",
            "amount": "4500",
        }, out_path)

        assert out_path.exists()
        # Verify no tokens remain
        remaining = extract_tokens(out_path)
        assert len(remaining) == 0

    def test_fill_result_has_counts(self, tmp_path):
        from doc_utils import fill_template
        tmpl_path = tmp_path / "tmpl.docx"
        out_path = tmp_path / "out.docx"
        doc = Document()
        doc.add_paragraph("{{a}} and {{b}}")
        doc.save(str(tmpl_path))

        result = fill_template(tmpl_path, {"a": "1", "b": "2"}, out_path)
        assert isinstance(result, dict)

    def test_fill_missing_token_stays(self, tmp_path):
        from doc_utils import fill_template, extract_tokens
        tmpl_path = tmp_path / "tmpl.docx"
        out_path = tmp_path / "out.docx"
        doc = Document()
        doc.add_paragraph("{{filled}} {{unfilled}}")
        doc.save(str(tmpl_path))

        fill_template(tmpl_path, {"filled": "yes"}, out_path)
        remaining = extract_tokens(out_path)
        assert "unfilled" in remaining
        assert "filled" not in remaining


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestCreateTemplateFromDoc:
    def test_create_template_replaces_text(self, tmp_path):
        from doc_utils import create_template_from_doc, extract_tokens
        src_path = tmp_path / "src.docx"
        tmpl_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("Dear Acme Corp,")
        doc.add_paragraph("Amount: 4500 SGD")
        doc.save(str(src_path))

        create_template_from_doc(src_path, {
            "Acme Corp": "client_name",
            "4500": "amount",
            "SGD": "currency",
        }, tmpl_path)

        tokens = extract_tokens(tmpl_path)
        assert "client_name" in tokens
        assert "amount" in tokens
        assert "currency" in tokens


class TestRegisterTemplate:
    def test_register_creates_index(self, tmp_path):
        from doc_utils import register_template
        index_path = tmp_path / "index.yaml"
        register_template(
            name="Test NDA",
            file="templates/test_nda.docx",
            tokens=["client_name", "date"],
            category="legal",
            index_path=str(index_path),
        )
        assert index_path.exists()
        import yaml
        with open(index_path) as f:
            data = yaml.safe_load(f)
        assert "templates" in data
        assert len(data["templates"]) == 1
        assert data["templates"][0]["name"] == "Test NDA"
        assert "client_name" in data["templates"][0]["tokens"]

    def test_register_appends_to_existing(self, tmp_path):
        from doc_utils import register_template
        index_path = tmp_path / "index.yaml"
        # Create initial
        register_template("First", "f.docx", ["a"], "legal", str(index_path))
        # Add second
        register_template("Second", "s.docx", ["b"], "finance", str(index_path))
        import yaml
        with open(index_path) as f:
            data = yaml.safe_load(f)
        assert len(data["templates"]) == 2